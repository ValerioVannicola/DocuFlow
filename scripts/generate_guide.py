"""Generate the DocuFlow User Guide as a Word document."""
from __future__ import annotations

import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(9)
code_font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
code_fmt = code_style.paragraph_format
code_fmt.space_before = Pt(4)
code_fmt.space_after = Pt(4)
code_fmt.left_indent = Inches(0.3)


def add_code(text):
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')


def add_param_table(params):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light List Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Type'
    hdr[2].text = 'Default'
    for name, typ, default in params:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = typ
        row[2].text = default


# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('DocuFlow User Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Complete Reference for the DocuFlow Document Processing Library', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Version 0.1.0')
doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS placeholder
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Installation',
    '3. Quick Start',
    '4. Architecture Overview',
    '5. Document Models',
    '6. Pipeline',
    '7. Parsing',
    '8. Extraction Engines',
    '9. Templates',
    '10. Validation',
    '11. Review',
    '12. Privacy & Anonymization',
    '13. Storage',
    '14. Observability',
    '15. CLI',
    '16. Error Handling',
    '17. API Reference',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'DocuFlow is a Python library for building production-grade document processing pipelines. '
    'It converts business documents (PDFs, scanned images) into validated, auditable structured '
    'data using AI-powered extraction with OCR and LLM support.'
)
doc.add_paragraph('Key capabilities:')
for cap in [
    'Extract structured data from PDFs using Pydantic schemas',
    'Multiple extraction strategies: text-based, vision-based, and hybrid',
    'Single and multi-agent extraction with decider consensus',
    'Tesseract OCR integration for scanned documents',
    'Privacy-first: PII detection, anonymization, and redaction via Presidio',
    'Configurable review rules and LLM-powered reviewers',
    'Human correction tracking with full audit trail',
    'Local file storage with JSON persistence',
    'Async-first with sync convenience wrappers',
]:
    doc.add_paragraph(cap, style='List Bullet')

# ============================================================
# 2. INSTALLATION
# ============================================================
doc.add_heading('2. Installation', level=1)
doc.add_paragraph('Install the full library with all optional dependencies:')
add_code('pip install docuflow[all]')
doc.add_paragraph('Or install only what you need:')
add_code('''pip install docuflow[pdf]      # PDF parsing with PyMuPDF
pip install docuflow[ocr]      # OCR with Tesseract
pip install docuflow[llm]      # LLM extraction via litellm
pip install docuflow[privacy]  # Privacy/anonymization via Presidio''')
doc.add_paragraph('Requirements: Python >= 3.11')

doc.add_heading('Optional Dependencies', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Group'
hdr[1].text = 'Packages'
hdr[2].text = 'Purpose'
for group, pkgs, purpose in [
    ('pdf', 'pymupdf', 'PDF text extraction and rendering'),
    ('ocr', 'pytesseract, Pillow', 'OCR with Tesseract'),
    ('llm', 'litellm', 'LLM calls (OpenAI, Anthropic, Gemini)'),
    ('privacy', 'presidio-analyzer, presidio-anonymizer', 'PII detection and anonymization'),
    ('all', 'All of the above', 'Everything'),
]:
    row = table.add_row().cells
    row[0].text = group
    row[1].text = pkgs
    row[2].text = purpose

# ============================================================
# 3. QUICK START
# ============================================================
doc.add_heading('3. Quick Start', level=1)

doc.add_heading('One-liner extraction', level=2)
add_code('''from pydantic import BaseModel
from docuflow import extract

class Invoice(BaseModel):
    supplier_name: str
    invoice_number: str
    total: float

result = extract("invoice.pdf", schema=Invoice)
print(result.data)
print(result.fields["total"].evidence)''')

doc.add_heading('Using a built-in template', level=2)
add_code('''from docuflow import extract
from docuflow.templates import load_template

Invoice = load_template("invoice")
result = extract("invoice.pdf", schema=Invoice)''')

doc.add_heading('Reusable pipeline', level=2)
add_code('''from docuflow import DocumentPipeline

pipeline = DocumentPipeline(
    parser="tesseract",
    model="openai/gpt-4o",
    storage="local",
)
result = pipeline.run_sync("invoice.pdf", schema=Invoice)''')

# ============================================================
# 4. ARCHITECTURE OVERVIEW
# ============================================================
doc.add_heading('4. Architecture Overview', level=1)
doc.add_paragraph(
    'DocuFlow processes documents through a configurable pipeline of steps. '
    'Each step receives a PipelineState, performs its work, and passes the '
    'updated state to the next step.'
)

doc.add_heading('Pipeline Flow', level=2)
doc.add_paragraph('The standard pipeline flow is:')
add_code('''Ingest -> Parse -> [Anonymize] -> Extract -> [Validate] -> [Review] -> [Store]''')

doc.add_paragraph('Steps in brackets are optional. The extraction step has three variants:')
table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Extraction Type'
hdr[1].text = 'Pipeline'
hdr[2].text = 'Description'
for etype, pipeline_str, desc in [
    ('text', 'Ingest -> Parse -> Extract', 'Parser extracts text, LLM reads text'),
    ('vision', 'Ingest -> ExtractVision', 'Pages rendered as images, sent to vision LLM'),
    ('hybrid', 'Ingest -> ExtractHybrid', 'Vision + text agents in parallel, vision decider'),
]:
    row = table.add_row().cells
    row[0].text = etype
    row[1].text = pipeline_str
    row[2].text = desc

doc.add_heading('Three Ways to Use DocuFlow', level=2)
doc.add_paragraph('1. One-liner function - simplest, least control:')
add_code('result = extract("file.pdf", schema=Invoice)')
doc.add_paragraph('2. DocumentPipeline - configurable, reusable:')
add_code('''pipeline = DocumentPipeline(parser="tesseract", model="openai/gpt-4o")
result = pipeline.run_sync("file.pdf", schema=Invoice)''')
doc.add_paragraph('3. Manual Pipeline - full control over step sequence:')
add_code('''from docuflow.workflow import Pipeline, Ingest, Parse, Extract, Store

pipeline = Pipeline([
    Ingest(path="file.pdf"),
    Parse(parser="tesseract"),
    Extract(schema=Invoice, llm=my_llm, mode="multi", n_instances=3),
    Store(storage=LocalDocumentStore("./output")),
])
result = pipeline.run_sync()''')

# ============================================================
# 5. DOCUMENT MODELS
# ============================================================
doc.add_heading('5. Document Models', level=1)
doc.add_paragraph('The core data structures that flow through the pipeline.')

doc.add_heading('Document', level=2)
doc.add_paragraph('Represents a file and all derived content: pages, blocks, text.')
add_param_table([
    ('id', 'str', 'UUID'),
    ('metadata', 'DocumentMetadata', '(required)'),
    ('pages', 'list[Page]', '[]'),
    ('raw_text', 'str', '""'),
    ('status', 'str', '"ingested"'),
])
add_code('''from docuflow.documents.models import Document
doc = await Document.from_file("invoice.pdf")
doc = Document.from_file_sync("invoice.pdf")''')

doc.add_heading('Page', level=2)
doc.add_paragraph('One page of the document with text and layout blocks.')
add_param_table([
    ('page_number', 'int', '(required)'),
    ('width', 'float | None', 'None'),
    ('height', 'float | None', 'None'),
    ('blocks', 'list[Block]', '[]'),
    ('text', 'str', '""'),
    ('image_path', 'str | None', 'None'),
])

doc.add_heading('Block', level=2)
doc.add_paragraph('A text or image element on a page with position and optional confidence.')
add_param_table([
    ('block_id', 'str', '(required)'),
    ('block_type', 'BlockType', 'TEXT'),
    ('text', 'str', '""'),
    ('bbox', 'BoundingBox | None', 'None'),
    ('confidence', 'float | None', 'None'),
])
doc.add_paragraph('BlockType values: TEXT, TITLE, TABLE, IMAGE, HEADER, FOOTER, LIST_ITEM, FORMULA, PARAGRAPH')

doc.add_heading('Evidence', level=2)
doc.add_paragraph('Links an extracted value back to its source location in the document.')
add_param_table([
    ('document_id', 'str', '(required)'),
    ('page_number', 'int', '(required)'),
    ('text', 'str', '(required)'),
    ('bbox', 'BoundingBox | None', 'None'),
    ('block_id', 'str | None', 'None'),
    ('confidence', 'float | None', 'None'),
])

doc.add_heading('ExtractionResult', level=2)
doc.add_paragraph('The complete output of the extraction pipeline.')
add_param_table([
    ('document_id', 'str', '(required)'),
    ('schema_name', 'str', '(required)'),
    ('data', 'dict', '{}'),
    ('fields', 'dict[str, ExtractedField]', '{}'),
    ('confidence', 'float', '0.0'),
    ('needs_review', 'bool', 'False'),
    ('review_status', 'str', '"pending"'),
    ('review_reasons', 'list[str]', '[]'),
    ('review_verdicts', 'list[ReviewVerdict]', '[]'),
    ('corrections', 'list[FieldCorrection]', '[]'),
    ('validation_errors', 'list[dict]', '[]'),
    ('trace_id', 'str', '""'),
    ('model_name', 'str', '""'),
])

doc.add_heading('ExtractedField', level=2)
doc.add_paragraph('A single extracted field with confidence, evidence, and correction tracking.')
add_param_table([
    ('value', 'Any', 'None'),
    ('original_value', 'Any', 'None'),
    ('corrected', 'bool', 'False'),
    ('confidence', 'float', '0.0'),
    ('evidence', 'list[Evidence]', '[]'),
    ('validation_status', 'str', '"pending"'),
    ('errors', 'list[str]', '[]'),
])

# ============================================================
# 6. PIPELINE
# ============================================================
doc.add_heading('6. Pipeline', level=1)

doc.add_heading('DocumentPipeline', level=2)
doc.add_paragraph('The main entry point for configuring and running document processing.')
add_param_table([
    ('parser', 'str | Parser', '"pymupdf"'),
    ('model', 'str', '"openai/gpt-4o"'),
    ('storage', 'str | Storage | None', 'None'),
    ('validators', 'list | None', 'None'),
    ('review_rules', 'list | None', 'None'),
    ('privacy', 'PrivacyPolicy | None', 'None'),
    ('extraction_mode', 'str', '"single"'),
    ('extraction_type', 'str', '"text"'),
    ('n_instances', 'int', '5'),
    ('temperatures', 'list[float] | None', 'None'),
    ('vision_dpi', 'int', '200'),
])

doc.add_heading('Text extraction pipeline', level=3)
add_code('''pipeline = DocumentPipeline(
    parser="pymupdf",          # or "tesseract" for scanned docs
    model="openai/gpt-4o",
    extraction_mode="single",  # or "multi" for multi-agent
    n_instances=3,             # number of parallel agents (multi mode)
)
result = pipeline.run_sync("invoice.pdf", schema=Invoice)''')

doc.add_heading('Vision extraction pipeline', level=3)
add_code('''pipeline = DocumentPipeline(
    parser=None,               # no parser needed
    extraction_type="vision",  # send page images to vision LLM
    model="openai/gpt-4o",
    extraction_mode="multi",
    n_instances=3,
)
result = pipeline.run_sync("scanned.pdf", schema=Invoice)''')

doc.add_heading('Hybrid extraction pipeline', level=3)
add_code('''pipeline = DocumentPipeline(
    parser=None,
    extraction_type="hybrid",  # vision + text agents in parallel
    model="openai/gpt-4o",
    n_instances=2,             # 2 vision + 2 text + 1 decider = 5 LLM calls
)
result = pipeline.run_sync("complex.pdf", schema=Invoice)''')

doc.add_heading('Full-featured pipeline', level=3)
add_code('''from docuflow import DocumentPipeline, PrivacyPolicy
from docuflow.privacy import PresidioProvider
from docuflow.review import OverallConfidenceBelow, FieldMissing, LLMReviewer
from docuflow.validation import RequiredFields, EvidenceRequired

pipeline = DocumentPipeline(
    parser="tesseract",
    model="openai/gpt-4o",
    storage="local",
    privacy=PrivacyPolicy(
        provider=PresidioProvider(),
        mode="pseudonymize",
    ),
    validators=[
        RequiredFields(["supplier_name", "total"]),
        EvidenceRequired(),
    ],
    review_rules=[
        OverallConfidenceBelow(0.7),
        FieldMissing(["total", "invoice_number"]),
    ],
    extraction_mode="multi",
    n_instances=3,
)
result = pipeline.run_sync("invoice.pdf", schema=Invoice)''')

# ============================================================
# 7. PARSING
# ============================================================
doc.add_heading('7. Parsing', level=1)
doc.add_paragraph('Parsing converts raw PDF files into structured Document objects with text and layout.')

doc.add_heading('PyMuPDFParser', level=2)
doc.add_paragraph(
    'Extracts embedded text from digital PDFs. Fast, no OCR needed. '
    'Does not work on scanned documents (returns empty text).'
)
add_code('''from docuflow.parsing.pymupdf import PyMuPDFParser

parser = PyMuPDFParser()
document = await parser.parse(document)
# document.pages[0].text -> extracted text
# document.pages[0].blocks -> text blocks with bounding boxes''')

doc.add_heading('TesseractParser', level=2)
doc.add_paragraph(
    'Renders PDF pages to images and runs Tesseract OCR. Works on scanned documents. '
    'Produces blocks with per-word confidence scores and bounding boxes.'
)
add_param_table([
    ('languages', 'list[str]', '["eng"]'),
    ('dpi', 'int', '200'),
    ('preprocess_steps', 'list[str] | None', 'None'),
])
add_code('''from docuflow.parsing.tesseract_parser import TesseractParser

parser = TesseractParser(languages=["eng", "ita"], dpi=300)
document = await parser.parse(document)
# document.pages[0].blocks[0].confidence -> OCR confidence (0-1)''')

doc.add_heading('Comparison', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'PyMuPDF'
hdr[2].text = 'Tesseract'
hdr[3].text = ''
for feat, pymupdf, tess, _ in [
    ('Best for', 'Digital PDFs', 'Scanned documents', ''),
    ('Speed', 'Fast (< 100ms)', 'Slower (1-5s/page)', ''),
    ('Bounding boxes', 'Yes', 'Yes', ''),
    ('Block confidence', 'No (None)', 'Yes (0-1 per word)', ''),
    ('Requires', 'pymupdf', 'pytesseract + Pillow', ''),
]:
    row = table.add_row().cells
    row[0].text = feat
    row[1].text = pymupdf
    row[2].text = tess

# ============================================================
# 8. EXTRACTION ENGINES
# ============================================================
doc.add_heading('8. Extraction Engines', level=1)

doc.add_heading('Extraction Modes', level=2)
doc.add_paragraph('Every extraction engine supports two modes:')
table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Mode'
hdr[1].text = 'Description'
hdr[2].text = 'LLM Calls'
for mode, desc, calls in [
    ('single', 'One LLM call', '1'),
    ('multi', 'N parallel calls at varied temperatures, then a decider', 'N + 1'),
]:
    row = table.add_row().cells
    row[0].text = mode
    row[1].text = desc
    row[2].text = calls

doc.add_heading('ExtractionEngine (text-based)', level=2)
doc.add_paragraph('Sends parsed document text to the LLM. Requires a Parse step before it.')
add_code('''from docuflow.extraction.engine import ExtractionEngine
from docuflow.extraction.llm.litellm_adapter import LiteLLMAdapter

llm = LiteLLMAdapter(model="openai/gpt-4o")
engine = ExtractionEngine(llm=llm)
result = await engine.extract(document, schema=Invoice, mode="single")
result = await engine.extract(document, schema=Invoice, mode="multi", n_instances=3)''')

doc.add_heading('VisionExtractionEngine', level=2)
doc.add_paragraph(
    'Renders PDF pages to images and sends them to a vision LLM. '
    'Automatically runs Tesseract OCR on the images for evidence grounding '
    '(bounding boxes, block confidence). No Parse step needed.'
)
add_code('''from docuflow.extraction.engine import VisionExtractionEngine

engine = VisionExtractionEngine(llm=llm, dpi=200)
result = await engine.extract(document, schema=Invoice, mode="single")''')

doc.add_heading('HybridExtractionEngine', level=2)
doc.add_paragraph(
    'Runs vision and text agents in parallel for maximum diversity. '
    'N vision agents read page images, N text agents read OCR markdown. '
    'A vision decider reviews all candidates against the actual page images.'
)
add_code('''from docuflow.extraction.engine import HybridExtractionEngine

engine = HybridExtractionEngine(llm=llm, dpi=200)
result = await engine.extract(document, schema=Invoice, n_instances=2)
# 2 vision + 2 text + 1 decider = 5 LLM calls''')

doc.add_heading('Confidence Scoring', level=2)
doc.add_paragraph('Confidence is computed per field from two sources:')
for item in [
    'LLM self-reported confidence (0-1) as returned in the evidence',
    'OCR confidence from Tesseract (when OCR was used), averaged with LLM confidence',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('JSON Reliability', level=2)
doc.add_paragraph('All LLM calls enforce reliable JSON output:')
for item in [
    'JSON mode (response_format={"type": "json_object"}) on every call',
    'Concrete output example with actual field names in the prompt',
    'Markdown fence stripping if LLM wraps JSON in code blocks',
    'Automatic retry with repair prompt if JSON parsing fails',
]:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 9. TEMPLATES
# ============================================================
doc.add_heading('9. Templates', level=1)
doc.add_paragraph(
    'Templates are YAML files that define extraction schemas. '
    'They get dynamically converted to Pydantic models at runtime.'
)

doc.add_heading('Built-in Templates', level=2)
for name, desc in [
    ('invoice', 'supplier, number, dates, totals, VAT, line items'),
    ('contract', 'type, parties, dates, terms, liability, governing law'),
    ('receipt', 'merchant, date, total, payment method, items'),
]:
    doc.add_paragraph(f'{name} - {desc}', style='List Bullet')

doc.add_heading('Using Templates', level=2)
add_code('''from docuflow.templates import load_template, list_templates

# List available templates
for t in list_templates():
    print(f"{t.name} ({t.source}): {t.description}")

# Load a template
Invoice = load_template("invoice")
result = extract("invoice.pdf", schema=Invoice)''')

doc.add_heading('Creating Custom Templates', level=2)
add_code('''# Save as ./docuflow_templates/my_schema.yaml
name: my_schema
version: "1.0"
description: "My custom extraction schema"
fields:
  company_name:
    type: str
    required: true
    description: "Company name"
  revenue:
    type: float
    required: true
  fiscal_year:
    type: date
    required: false''')
doc.add_paragraph(
    'Template discovery order: ./docuflow_templates/ > ~/.docuflow/templates/ > built-in. '
    'First match wins.'
)

# ============================================================
# 10. VALIDATION
# ============================================================
doc.add_heading('10. Validation', level=1)
doc.add_paragraph('Validators check extraction results against rules and update field statuses.')

doc.add_heading('Built-in Validators', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Validator'
hdr[1].text = 'Parameters'
hdr[2].text = 'What it checks'
for name, params, desc in [
    ('RequiredFields', 'fields: list[str]', 'Fields must be present and non-None'),
    ('TypeValidation', '(none)', 'Warns on empty string values'),
    ('EvidenceRequired', 'fields: list[str] | None', 'Fields must have evidence'),
    ('CustomRule', 'name, fn', 'User-defined validation function'),
]:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = params
    row[2].text = desc

doc.add_heading('Example', level=2)
add_code('''from docuflow.validation import RequiredFields, EvidenceRequired

pipeline = DocumentPipeline(
    validators=[
        RequiredFields(["supplier_name", "total", "invoice_number"]),
        EvidenceRequired(["total"]),
    ],
)''')

# ============================================================
# 11. REVIEW
# ============================================================
doc.add_heading('11. Review', level=1)
doc.add_paragraph(
    'The Review step checks extraction results against rules and LLM reviewers. '
    'If any rule triggers, the document is flagged with needs_review=True.'
)

doc.add_heading('Review Rules', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Rule'
hdr[1].text = 'Parameters'
hdr[2].text = 'When it flags'
for name, params, desc in [
    ('OverallConfidenceBelow', 'threshold=0.7', 'Average confidence below threshold'),
    ('FieldConfidenceBelow', 'fields: dict[str, float]', 'Per-field confidence thresholds'),
    ('AnyFieldConfidenceBelow', 'threshold=0.6', 'Any single field below threshold'),
    ('HasValidationErrors', '(none)', 'Validation step found errors'),
    ('FieldMissing', 'fields: list[str]', 'Critical fields are None or absent'),
    ('NoEvidence', 'fields: list[str] | None', 'Fields have no supporting evidence'),
]:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = params
    row[2].text = desc

doc.add_heading('LLM Reviewer', level=2)
doc.add_paragraph(
    'LLM reviewers are prompt-driven agents that review extractions. '
    'Each reviewer has a name, a custom prompt, and returns a structured '
    'ReviewVerdict with "Approved" or "Not Approved" and reasoning.'
)
add_code('''from docuflow.review import LLMReviewer, OverallConfidenceBelow
from docuflow.extraction.llm.litellm_adapter import LiteLLMAdapter

llm = LiteLLMAdapter(model="openai/gpt-4o")

auditor = LLMReviewer(
    name="financial_auditor",
    prompt="Check if extracted totals, VAT, and line items are mathematically consistent.",
    llm=llm,
)

compliance = LLMReviewer(
    name="compliance_check",
    prompt="Check if any extracted field contains PII that should not be there.",
    llm=llm,
)

pipeline = DocumentPipeline(
    review_rules=[
        OverallConfidenceBelow(0.7),
        auditor,
        compliance,
    ],
)''')

doc.add_heading('Review Verdicts', level=2)
doc.add_paragraph('Each LLM reviewer produces a ReviewVerdict stored on the result:')
add_code('''for v in result.review_verdicts:
    print(f"{v.reviewer}: {v.verdict} - {v.reasoning}")
# financial_auditor: Approved - Math checks out
# compliance_check: Not Approved - PII found in supplier field''')

doc.add_heading('Human Corrections', level=2)
doc.add_paragraph('After review, humans can correct fields with a full audit trail:')
add_code('''# Correct a field
result.correct_field(
    "total", 1235.00,
    corrected_by="john",
    reason="OCR misread 5 as 6"
)

# The original value is preserved
print(result.fields["total"].value)           # 1235.00
print(result.fields["total"].original_value)  # 1234.56
print(result.fields["total"].corrected)       # True

# Full correction history
for c in result.corrections:
    print(f"{c.field_name}: {c.old_value} -> {c.new_value} by {c.corrected_by}")''')

doc.add_heading('Approve / Reject', level=2)
add_code('''# Approve after corrections
result.approve(approved_by="john")
print(result.review_status)  # "approved"

# Or reject
result.reject(rejected_by="john", reason="wrong document type")
print(result.review_status)   # "rejected"
print(result.rejection_reason)  # "wrong document type"''')

# ============================================================
# 12. PRIVACY & ANONYMIZATION
# ============================================================
doc.add_heading('12. Privacy & Anonymization', level=1)
doc.add_paragraph(
    'The privacy module detects and anonymizes PII before it reaches the LLM. '
    'It runs as a pipeline step between Parse and Extract.'
)

doc.add_heading('Anonymization Modes', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Mode'
hdr[1].text = 'Example'
hdr[2].text = 'Use case'
for mode, example, use in [
    ('redact', 'Mario Rossi -> [REDACTED]', 'Maximum privacy'),
    ('mask', 'Mario Rossi -> M**** R****', 'Partial masking for review'),
    ('pseudonymize', 'Mario Rossi -> PERSON_001', 'LLM workflows (reversible)'),
    ('hash', 'Mario Rossi -> a3f2...', 'Duplicate detection'),
]:
    row = table.add_row().cells
    row[0].text = mode
    row[1].text = example
    row[2].text = use

doc.add_heading('Basic Usage', level=2)
add_code('''from docuflow import DocumentPipeline, PrivacyPolicy
from docuflow.privacy import PresidioProvider

pipeline = DocumentPipeline(
    parser="tesseract",
    privacy=PrivacyPolicy(
        provider=PresidioProvider(),
        mode="pseudonymize",
        reversible=True,
        fail_closed=True,
        entities=["PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "IBAN_CODE"],
    ),
)
result = pipeline.run_sync("claim.pdf", schema=ClaimSchema)''')

doc.add_heading('PrivacyPolicy Parameters', level=2)
add_param_table([
    ('anonymize_before_llm', 'bool', 'True'),
    ('mode', 'AnonymizationMode', '"pseudonymize"'),
    ('reversible', 'bool', 'True'),
    ('provider', 'PrivacyProvider', 'None'),
    ('entities', 'list[str]', '["PERSON", "EMAIL_ADDRESS", ...]'),
    ('fail_closed', 'bool', 'True'),
    ('score_threshold', 'float', '0.35'),
    ('log_scrubbing', 'bool', 'True'),
    ('mapping_store', 'MappingStore | None', 'None'),
])

doc.add_heading('Reversible Pseudonymization', level=2)
add_code('''from docuflow.privacy import Anonymizer, PresidioProvider
from docuflow.privacy.mapping_store import LocalMappingStore

anonymizer = Anonymizer(PrivacyPolicy(
    provider=PresidioProvider(),
    mode="pseudonymize",
    reversible=True,
    mapping_store=LocalMappingStore("./mappings"),
))

anon = await anonymizer.anonymize_text("John Doe sent email john@example.com")
print(anon.text)  # "PERSON_001 sent email EMAIL_ADDRESS_001"

restored = await anonymizer.restore_text(anon.text, anon.mapping_id)
print(restored)    # "John Doe sent email john@example.com"''')

doc.add_heading('Image Redaction', level=2)
add_code('''from docuflow.privacy.image_redaction import ImageRedactor

redactor = ImageRedactor(provider=PresidioProvider())
redacted_image, findings = await redactor.redact_page_image(page_image)
# Black rectangles drawn over PII regions using OCR bounding boxes''')

doc.add_heading('Trace Scrubbing', level=2)
add_code('''from docuflow.privacy.scrubber import TraceScrubber

scrubber = TraceScrubber(provider=PresidioProvider())
clean_trace = await scrubber.scrub_trace(result.trace)
# PII replaced with [SCRUBBED] in all trace event metadata''')

# ============================================================
# 13. STORAGE
# ============================================================
doc.add_heading('13. Storage', level=1)
doc.add_paragraph('Storage persists documents, extraction results, and traces to disk.')

doc.add_heading('LocalDocumentStore', level=2)
add_code('''from docuflow.storage.local import LocalDocumentStore

store = LocalDocumentStore("./docuflow_output")

# Save
await store.save_document(document)
await store.save_result(extraction_result)
await store.save_trace(trace)

# Load
result = await store.load_result("document-id")
document = await store.load_document("document-id")''')

doc.add_paragraph('Files saved per document:')
for item in [
    'original.pdf - copy of the source file',
    'document.json - parsed Document (pages, blocks, text)',
    'extraction.json - ExtractionResult (fields, values, evidence, corrections)',
    'trace.json - processing trace (events, timing)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'On pipeline failure with storage="local", partial state is automatically saved '
    'so you can inspect what went wrong.'
)

# ============================================================
# 14. OBSERVABILITY
# ============================================================
doc.add_heading('14. Observability', level=1)
doc.add_paragraph('Every pipeline step records trace events with timing information.')
add_code('''# Access the trace from PipelineResult
pipeline = Pipeline([Ingest(path="file.pdf"), Parse(), Extract(...)])
result = pipeline.run_sync()

for event in result.trace.events:
    print(f"{event.event_type}: {event.step_name} ({event.duration_ms:.0f}ms)")
# ingest: ingest (15ms)
# parse: parse (234ms)
# extract: extraction (1850ms)''')

# ============================================================
# 15. CLI
# ============================================================
doc.add_heading('15. CLI', level=1)
doc.add_paragraph('DocuFlow includes a command-line interface for common tasks.')

add_code('''# Extract data from a document
docuflow extract invoice.pdf --schema invoice --model openai/gpt-4o

# Save output to file
docuflow extract invoice.pdf --schema invoice --output result.json

# List available templates
docuflow templates list

# Show a template definition
docuflow templates show invoice

# Copy a template for customization
docuflow templates init invoice --dir ./my_templates''')

# ============================================================
# 16. ERROR HANDLING
# ============================================================
doc.add_heading('16. Error Handling', level=1)
doc.add_paragraph('All errors inherit from DocuflowError for easy catching.')
table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Exception'
hdr[1].text = 'When raised'
for exc, when in [
    ('DocuflowError', 'Base class for all docuflow errors'),
    ('UnsupportedFileTypeError', 'Unknown file extension'),
    ('ParsingError', 'PDF parsing or file access failure'),
    ('OCRError', 'Tesseract OCR failure'),
    ('SchemaExtractionError', 'LLM call or JSON parse failure'),
    ('ValidationError', 'Field validation failure'),
    ('StorageError', 'Storage read/write failure'),
    ('WorkflowError', 'Pipeline step failure (carries result)'),
    ('PrivacyError', 'Privacy operation failure'),
    ('AnonymizationError', 'Anonymization failure'),
]:
    row = table.add_row().cells
    row[0].text = exc
    row[1].text = when

doc.add_heading('Handling Pipeline Failures', level=2)
add_code('''from docuflow.errors import WorkflowError

try:
    result = pipeline.run_sync("invoice.pdf", schema=Invoice)
except WorkflowError as e:
    print(e)                           # error message
    print(e.result.errors)             # list of step errors
    print(e.result.state.current_step) # which step failed
    print(e.result.trace.events)       # trace up to failure
    print(e.result.state.document)     # document if ingestion succeeded''')

# ============================================================
# 17. API REFERENCE
# ============================================================
doc.add_heading('17. API Reference Summary', level=1)

doc.add_heading('Top-level imports', level=2)
add_code('''from docuflow import extract, DocumentPipeline, Pipeline, PrivacyPolicy

from docuflow.templates import load_template, list_templates
from docuflow.validation import RequiredFields, EvidenceRequired
from docuflow.review import (
    OverallConfidenceBelow, FieldConfidenceBelow, AnyFieldConfidenceBelow,
    HasValidationErrors, FieldMissing, NoEvidence, LLMReviewer,
)
from docuflow.privacy import PrivacyPolicy, Anonymizer, PresidioProvider
from docuflow.storage.local import LocalDocumentStore
from docuflow.extraction.llm.litellm_adapter import LiteLLMAdapter
from docuflow.parsing.pymupdf import PyMuPDFParser
from docuflow.parsing.tesseract_parser import TesseractParser
from docuflow.workflow import (
    Pipeline, Ingest, Parse, Extract, ExtractVision, ExtractHybrid,
    Anonymize, Validate, Review, Store,
)''')

# ============================================================
# SAVE
# ============================================================
output_path = os.path.expanduser(r"C:\Users\vvannicola\Downloads\DocuFlow_User_Guide.docx")
doc.save(output_path)
print(f"Guide saved to {output_path}")
