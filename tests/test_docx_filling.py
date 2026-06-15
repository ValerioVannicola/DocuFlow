"""Tests for DOCX form filling (content_controls and template strategies)."""
from __future__ import annotations

import asyncio
import io
from pathlib import Path
from typing import Any

import pytest
from pydantic import BaseModel, Field

# ---------------------------------------------------------------------------
# Helpers — build minimal in-memory DOCX files without touching disk
# ---------------------------------------------------------------------------

def _make_content_controls_docx(fields: dict[str, str]) -> bytes:
    """Build a DOCX with one plain-text SDT per field."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    body = doc.element.body

    for name, default in fields.items():
        sdt = OxmlElement("w:sdt")
        # Properties
        sdtPr = OxmlElement("w:sdtPr")
        tag_el = OxmlElement("w:tag")
        tag_el.set(qn("w:val"), name)
        alias_el = OxmlElement("w:alias")
        alias_el.set(qn("w:val"), name)
        text_el = OxmlElement("w:text")
        sdtPr.append(tag_el)
        sdtPr.append(alias_el)
        sdtPr.append(text_el)
        sdt.append(sdtPr)
        # Content
        sdtContent = OxmlElement("w:sdtContent")
        para = OxmlElement("w:p")
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = default
        run.append(t)
        para.append(run)
        sdtContent.append(para)
        sdt.append(sdtContent)
        body.append(sdt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_checkbox_docx(name: str, checked: bool = False) -> bytes:
    """Build a DOCX with a single modern w14:checkbox SDT."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from lxml import etree

    W14 = "http://schemas.microsoft.com/office/word/2010/wordml"

    doc = Document()
    body = doc.element.body

    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), name)
    sdtPr.append(tag_el)
    # Use lxml directly for the w14 namespace (not registered in python-docx)
    cb14 = etree.SubElement(sdtPr, f"{{{W14}}}checkbox")
    checked_el = etree.SubElement(cb14, f"{{{W14}}}checked")
    checked_el.set(f"{{{W14}}}val", "1" if checked else "0")
    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    para = OxmlElement("w:p")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "☑" if checked else "☐"
    run.append(t)
    para.append(run)
    sdtContent.append(para)
    sdt.append(sdtContent)
    body.append(sdt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_template_docx(text: str) -> bytes:
    """Build a DOCX whose first paragraph contains raw Jinja2 template text."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _write_tmp(tmp_path: Path, name: str, data: bytes) -> Path:
    p = tmp_path / name
    p.write_bytes(data)
    return p


# ---------------------------------------------------------------------------
# Model fixture
# ---------------------------------------------------------------------------

class ContractData(BaseModel):
    claimant_name: str = Field(alias="claimant.name")
    policy_number: str
    accepted_terms: bool = False


@pytest.fixture()
def contract() -> ContractData:
    return ContractData(
        **{"claimant.name": "Mario Rossi", "policy_number": "POL-42", "accepted_terms": True}
    )


# ---------------------------------------------------------------------------
# Inspector tests
# ---------------------------------------------------------------------------

def test_inspect_content_controls_finds_text_fields(tmp_path):
    p = _write_tmp(tmp_path, "form.docx", _make_content_controls_docx({"name": "", "city": ""}))
    from docuflow.filling.docx_inspector import inspect_content_controls

    fields = inspect_content_controls(p)
    names = {f.name for f in fields}
    assert {"name", "city"} == names
    assert all(f.field_type == "plainText" for f in fields)


def test_inspect_content_controls_checkbox(tmp_path):
    p = _write_tmp(tmp_path, "cb.docx", _make_checkbox_docx("accepted", checked=False))
    from docuflow.filling.docx_inspector import inspect_content_controls

    fields = inspect_content_controls(p)
    assert len(fields) == 1
    assert fields[0].field_type == "checkbox"
    assert fields[0].name == "accepted"


def test_inspect_template_vars(tmp_path):
    p = _write_tmp(tmp_path, "tpl.docx", _make_template_docx("Hello {{ first_name }}, your ref is {{ policy_number }}."))
    from docuflow.filling.docx_inspector import inspect_template_vars

    vars_ = inspect_template_vars(p)
    assert set(vars_) == {"first_name", "policy_number"}


def test_inspect_template_vars_empty_doc(tmp_path):
    p = _write_tmp(tmp_path, "empty.docx", _make_template_docx("No variables here."))
    from docuflow.filling.docx_inspector import inspect_template_vars

    assert inspect_template_vars(p) == []


# ---------------------------------------------------------------------------
# Strategy auto-detection
# ---------------------------------------------------------------------------

def test_auto_strategy_picks_content_controls(tmp_path):
    p = _write_tmp(tmp_path, "form.docx", _make_content_controls_docx({"policy_number": ""}))
    from docuflow.filling.docx_planner import _select_docx_strategy

    assert _select_docx_strategy(p, "auto") == "content_controls"


def test_auto_strategy_picks_template(tmp_path):
    p = _write_tmp(tmp_path, "tpl.docx", _make_template_docx("{{ policy_number }}"))
    from docuflow.filling.docx_planner import _select_docx_strategy

    assert _select_docx_strategy(p, "auto") == "template"


# ---------------------------------------------------------------------------
# Writer: content controls
# ---------------------------------------------------------------------------

def test_write_content_controls_text(tmp_path):
    src = _write_tmp(tmp_path, "src.docx", _make_content_controls_docx({"policy_number": "OLD", "claimant_name": ""}))
    out = tmp_path / "out.docx"

    from docuflow.filling.models import FilledField, FillingResult

    result = FillingResult(
        input_path=str(src),
        output_path=str(out),
        strategy="content_controls",
        fields={
            "policy_number": FilledField(
                field_name="policy_number",
                value="POL-42",
                formatted_value="POL-42",
                target_name="policy_number",
                status="filled",
            ),
        },
    )
    from docuflow.filling.docx_writer import write_content_controls

    write_content_controls(src, out, result)
    assert out.exists()

    from docuflow.filling.docx_inspector import inspect_content_controls

    fields = inspect_content_controls(out)
    by_name = {f.name: f for f in fields}
    assert by_name["policy_number"].current_value == "POL-42"


def test_write_content_controls_checkbox_true(tmp_path):
    src = _write_tmp(tmp_path, "src.docx", _make_checkbox_docx("accepted_terms", checked=False))
    out = tmp_path / "out.docx"

    from docuflow.filling.models import FilledField, FillingResult

    result = FillingResult(
        input_path=str(src),
        output_path=str(out),
        strategy="content_controls",
        fields={
            "accepted_terms": FilledField(
                field_name="accepted_terms",
                value=True,
                formatted_value=True,
                target_name="accepted_terms",
                status="filled",
            ),
        },
    )
    from docuflow.filling.docx_writer import write_content_controls

    write_content_controls(src, out, result)
    # Inspect the w14:checked val
    from docx import Document
    from docx.oxml.ns import qn

    W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
    doc = Document(str(out))
    for sdt in doc.element.iter(qn("w:sdt")):
        props = sdt.find(qn("w:sdtPr"))
        if props is None:
            continue
        cb14 = props.find(f"{{{W14}}}checkbox")
        if cb14 is not None:
            checked_el = cb14.find(f"{{{W14}}}checked")
            assert checked_el.get(f"{{{W14}}}val") == "1"


def test_write_template(tmp_path):
    src = _write_tmp(
        tmp_path, "tpl.docx",
        _make_template_docx("Name: {{ claimant_name }}. Policy: {{ policy_number }}.")
    )
    out = tmp_path / "out.docx"

    from docuflow.filling.models import FilledField, FillingResult

    result = FillingResult(
        input_path=str(src),
        output_path=str(out),
        strategy="template",
        fields={
            "claimant_name": FilledField(
                field_name="claimant_name",
                value="Mario Rossi",
                formatted_value="Mario Rossi",
                target_name="claimant_name",
                status="filled",
            ),
            "policy_number": FilledField(
                field_name="policy_number",
                value="POL-42",
                formatted_value="POL-42",
                target_name="policy_number",
                status="filled",
            ),
        },
    )
    from docuflow.filling.docx_writer import write_template

    write_template(src, out, result)
    assert out.exists()

    from docx import Document

    doc = Document(str(out))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "Mario Rossi" in text
    assert "POL-42" in text


# ---------------------------------------------------------------------------
# Full fill_docx_form end-to-end
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_fill_docx_form_content_controls_e2e(tmp_path, contract):
    src = _write_tmp(
        tmp_path, "form.docx",
        _make_content_controls_docx({"claimant_name": "", "policy_number": ""}),
    )
    out = tmp_path / "filled.docx"

    from docuflow.filling.api import fill_docx_form_async

    result = await fill_docx_form_async(
        str(src),
        {"claimant_name": "Mario Rossi", "policy_number": "POL-42"},
        output_path=str(out),
        strategy="content_controls",
    )
    assert result.success
    assert result.strategy == "content_controls"
    assert result.committed
    assert out.exists()
    assert not result.errors

    from docuflow.filling.docx_inspector import inspect_content_controls

    fields = inspect_content_controls(out)
    vals = {f.name: f.current_value for f in fields}
    assert vals["claimant_name"] == "Mario Rossi"
    assert vals["policy_number"] == "POL-42"


@pytest.mark.asyncio
async def test_fill_docx_form_template_e2e(tmp_path):
    src = _write_tmp(
        tmp_path, "tpl.docx",
        _make_template_docx("Dear {{ claimant_name }}, ref {{ policy_number }}."),
    )
    out = tmp_path / "filled.docx"

    from docuflow.filling.api import fill_docx_form_async

    result = await fill_docx_form_async(
        str(src),
        {"claimant_name": "Mario Rossi", "policy_number": "POL-42"},
        output_path=str(out),
        strategy="template",
    )
    assert result.success
    assert result.strategy == "template"
    assert result.committed
    assert out.exists()

    from docx import Document

    doc = Document(str(out))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "Mario Rossi" in text


@pytest.mark.asyncio
async def test_fill_docx_form_auto_strategy(tmp_path):
    src = _write_tmp(
        tmp_path, "form.docx",
        _make_content_controls_docx({"policy_number": ""}),
    )
    from docuflow.filling.api import fill_docx_form_async

    result = await fill_docx_form_async(str(src), {"policy_number": "POL-99"})
    assert result.strategy == "content_controls"
    assert result.success


# ---------------------------------------------------------------------------
# Review / approval workflow
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_fill_docx_review_defers_write(tmp_path):
    src = _write_tmp(
        tmp_path, "form.docx",
        _make_content_controls_docx({"policy_number": ""}),
    )
    out = tmp_path / "out.docx"

    from docuflow.filling.api import fill_docx_form_async

    result = await fill_docx_form_async(
        str(src), {"policy_number": "POL-42"}, output_path=str(out), review=True
    )
    assert not result.committed
    assert result.review_status == "pending"
    assert not out.exists()


@pytest.mark.asyncio
async def test_fill_docx_edit_approve_commit(tmp_path):
    src = _write_tmp(
        tmp_path, "form.docx",
        _make_content_controls_docx({"policy_number": ""}),
    )
    out = tmp_path / "out.docx"

    from docuflow.filling.api import commit_fill_async, fill_docx_form_async

    result = await fill_docx_form_async(
        str(src), {"policy_number": "POL-ORIG"}, output_path=str(out), review=True
    )
    result.edit_field("policy_number", value="POL-CORRECTED", corrected_by="test")
    result.approve(approved_by="test")
    await commit_fill_async(result)

    assert result.committed
    assert out.exists()
    assert result.corrections[0].old_value == "POL-ORIG"
    assert result.corrections[0].new_value == "POL-CORRECTED"

    from docuflow.filling.docx_inspector import inspect_content_controls

    fields = inspect_content_controls(out)
    vals = {f.name: f.current_value for f in fields}
    assert vals["policy_number"] == "POL-CORRECTED"


# ---------------------------------------------------------------------------
# Flatten
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_fill_docx_flatten_removes_sdts(tmp_path):
    src = _write_tmp(tmp_path, "form.docx", _make_content_controls_docx({"name": ""}))
    out = tmp_path / "flat.docx"

    from docuflow.filling.api import fill_docx_form_async

    await fill_docx_form_async(
        str(src), {"name": "Flat Test"}, output_path=str(out),
        strategy="content_controls", flatten=True
    )
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(out))
    sdts = list(doc.element.iter(qn("w:sdt")))
    assert len(sdts) == 0, "All SDTs should have been removed after flattening"


# ---------------------------------------------------------------------------
# Unmapped fields
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_fill_docx_unmapped_model_field_warns(tmp_path):
    src = _write_tmp(tmp_path, "form.docx", _make_content_controls_docx({"known": ""}))
    from docuflow.filling.api import fill_docx_form_async

    result = await fill_docx_form_async(
        str(src), {"known": "hello", "unknown_field": "ignored"},
    )
    assert result.success
    assert any("unknown_field" in w for w in result.warnings)
